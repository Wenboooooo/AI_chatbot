import faiss
import json
import numpy as np
from docx import Document
import openai
from openai import OpenAI
from autogen import AssistantAgent
from fastapi import FastAPI, WebSocket
from typing import Dict, List
import asyncio
import os
from dotenv import load_dotenv

# 加载环境变量
load_dotenv()

# 从环境变量获取 API 密钥
api_key = os.getenv("OPENAI_API_KEY")

# 使用 OpenAI 客户端
client = openai.OpenAI(api_key=api_key)

def extract_text_from_docx(file_path):
    doc = Document(file_path)
    return "\n".join([para.text.strip() for para in doc.paragraphs if para.text.strip()])

# 解析 .docx 和 .json 数据
docx_text = extract_text_from_docx("AI Doc.docx")

with open("jetbay-intro.json", "r", encoding="utf-8") as f:
    faq_data = json.load(f)

# 构建知识库数据
data = [{"text": docx_text, "source": "docx"}] + [{"text": item["question"] + " " + item["answer"], "source": "json"} for item in faq_data]

# 计算文本嵌入
def get_embedding(text):
    response = client.embeddings.create(input=[text], model="text-embedding-ada-002")  # Wrap input in a list
    return np.array(response.data[0].embedding) 

embeddings = np.array([get_embedding(item["text"]) for item in data])

# 创建 FAISS 索引
dimension = embeddings.shape[1]
faiss_index = faiss.IndexFlatL2(dimension)
faiss_index.add(embeddings)

# 存储索引
faiss.write_index(faiss_index, "jetbay_faiss.index")

# **FastAPI 应用**
app = FastAPI()

# **存储每个用户的聊天历史**
user_chat_histories: Dict[str, List[Dict[str, str]]] = {}

# **FAISS 检索代理**
class FAISSRetriever:
    def __init__(self, index_path, data):
        self.index = faiss.read_index(index_path)
        self.data = data
    
    def search(self, query, top_k=2):
        query_embedding = get_embedding(query).reshape(1, -1)
        D, I = self.index.search(query_embedding, top_k)
        results = [self.data[i]["text"] for i in I[0]]
        return results

# **创建嵌入获取函数**
def get_embedding(text):
    response = client.embeddings.create(input=[text], model="text-embedding-ada-002")  # 确保 input 是一个 list
    return np.array(response.data[0].embedding)

# **加载知识库数据**
def extract_text_from_docx(file_path):
    doc = Document(file_path)
    return "\n".join([para.text.strip() for para in doc.paragraphs if para.text.strip()])

docx_text = extract_text_from_docx("AI Doc.docx")

with open("jetbay-intro.json", "r", encoding="utf-8") as f:
    faq_data = json.load(f)

data = [{"text": docx_text, "source": "docx"}] + [{"text": item["question"] + " " + item["answer"], "source": "json"} for item in faq_data]

retriever = FAISSRetriever("jetbay_faiss.index", data)

# **创建 LLM 代理**
agent_prompt = """
You are an intelligent assistant in JetBay tasked with answering user queries based on the provided context.  
Use the context information to generate precise, relevant, and natural responses.  
Do not mention the existence of the context or that you are referencing it.  
"""

llm_agent = AssistantAgent(
    system_message=agent_prompt,
    name="llm",
    llm_config={
        "model": "gpt-4o",
        "api_key": api_key,
        "temperature": 0.7
    }
)


# **WebSocket 聊天 API**
# @app.websocket("/chat/{user_id}")
# async def websocket_endpoint(websocket: WebSocket, user_id: str):
#     await websocket.accept()
    
#     # 如果是新用户，则创建一个空的聊天历史记录
#     if user_id not in user_chat_histories:
#         user_chat_histories[user_id] = []
    
#     while True:
#         try:
#             # **接收前端用户消息**
#             user_query = await websocket.receive_text()

#             # **检索知识库**
#             retrieved_context = retriever.search(user_query, top_k=3)

#             # **存储用户输入**
#             user_chat_histories[user_id].append({"role": "user", "content": user_query})

#             # **构造对话历史**
#             conversation = user_chat_histories[user_id] + [{"role": "system", "content": f"Context: {retrieved_context}"}]

#             # **生成 AI 响应**
#             llm_response = llm_agent.generate_reply(conversation)

#             # **存储 AI 响应**
#             user_chat_histories[user_id].append({"role": "assistant", "content": llm_response})

#             # **发送 AI 响应**
#             await websocket.send_text(llm_response)

#         except Exception as e:
#             print(f"Error: {e}")
#             break

#     await websocket.close()



@app.websocket("/chat/{user_id}")
async def websocket_endpoint(websocket: WebSocket, user_id: str):
    await websocket.accept()
    
    if user_id not in user_chat_histories:
        user_chat_histories[user_id] = []

    while True:
        try:
            # **接收前端用户消息**
            user_query = await websocket.receive_text()

            # **检索知识库**
            retrieved_context = retriever.search(user_query, top_k=3)

            # **存储用户输入**
            user_chat_histories[user_id].append({"role": "user", "content": user_query})

            # **构造对话历史**
            conversation = user_chat_histories[user_id] + [{"role": "system", "content": f"Context: {retrieved_context}"}]

            # **向 OpenAI 发送流式请求**
            response_stream = client.chat.completions.create(
                model="gpt-4o",
                messages=conversation,
                temperature=0.7,
                stream=True  # 启用流式输出
            )

            # **存储 AI 响应**
            ai_response = ""

            # **逐步发送流式数据**
            async for chunk in response_stream:
                if chunk.choices[0].delta.content:
                    token = chunk.choices[0].delta.content
                    ai_response += token
                    await websocket.send_text(token)  # 逐步推送到客户端

            # **存储完整的 AI 响应**
            user_chat_histories[user_id].append({"role": "assistant", "content": ai_response})

        except Exception as e:
            print(f"Error: {e}")
            break

    await websocket.close()





if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)